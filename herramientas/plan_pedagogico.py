import fitz
import os
from mcp.server.fastmcp import FastMCP
from docx import Document
from openpyxl import load_workbook
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

mcp = FastMCP("eduforge-plan-pedagogico")

def leer_pdf(ruta: str) -> str:
    """Lee un archivo PDF y devuelve todo su texto"""
    documento = fitz.open(ruta)
    texto = ""
    for pagina in documento:
        texto += pagina.get_text()
    documento.close()
    return texto

def leer_plan_anual(ruta: str) -> str:
    """Lee el plan anual en formato .docx o .xlsx"""
    extension = os.path.splitext(ruta)[1].lower()
    if extension == ".docx":
        documento = Document(ruta)
        texto = ""
        for parrafo in documento.paragraphs:
            texto += parrafo.text + "\n"
        return texto
    elif extension == ".xlsx":
        wb = load_workbook(ruta)
        texto = ""
        for hoja in wb.worksheets:
            for fila in hoja.iter_rows():
                for celda in fila:
                    if celda.value:
                        texto += str(celda.value) + " "
                texto += "\n"
        return texto
    else:
        return "Formato no soportado. Usá .docx o .xlsx"

def generar_con_ia(prompt: str) -> str:
    """Llama a Claude para generar contenido — modo simulado por ahora"""
    respuesta_simulada = f"""
ESTRATEGIAS DEL DOCENTE (D/):
- Presenta los conceptos mediante técnica expositiva dialogada, 
  usando ejemplos del sector turístico y agrícola de Guanacaste.
- Modela el análisis mediante casos reales de empresas guanacastecas.
- Facilita discusión grupal conectando el tema con la experiencia 
  laboral del grupo nocturno.
- Promueve la reflexión usando analogías con actividades cotidianas 
  de la zona costera.

ESTRATEGIAS DEL ESTUDIANTE (E/):
- Construye mapa conceptual sobre los temas abordados.
- Analiza casos prácticos relacionados con su entorno laboral.
- Participa en discusión grupal aportando experiencias propias.

EVIDENCIAS:
Conocimiento: Mapa conceptual sobre los temas de la unidad.
Desempeño: Exposición grupal con ejemplos de Guanacaste.
Producto: Cuadro comparativo elaborado en clase.
"""
    return respuesta_simulada

@mcp.tool()
def generar_plan_pedagogico(
    ruta_programa: str,
    ruta_plan_anual: str,
    subarea: str,
    nivel: str,
    modalidad: str,
    nombre_docente: str,
    nombre_institucion: str,
    caracteristicas_grupo: str
) -> str:
    """Genera un plan de práctica pedagógica MEP completo"""
    texto_programa = leer_pdf(ruta_programa)
    texto_plan_anual = leer_plan_anual(ruta_plan_anual)
    prompt = f"""
Sos un docente experto en educación técnica costarricense.
Generá estrategias de mediación pedagógica para un plan de 
práctica pedagógica del MEP con estas características:

Subárea: {subarea}
Nivel: {nivel}
Modalidad: {modalidad}
Institución: {nombre_institucion}
Docente: {nombre_docente}
Características del grupo: {caracteristicas_grupo}

Contenido del programa de estudio:
{texto_programa[:3000]}

Plan anual:
{texto_plan_anual[:1000]}

Generá:
1. Estrategias del DOCENTE (D/) — mínimo 4, contextualizadas 
   para Guanacaste, fomentando la Guanacastequidad
2. Estrategias del ESTUDIANTE (E/) — basadas en indicadores de logro
3. Evidencias — Conocimiento, Desempeño y Producto
"""
    return generar_con_ia(prompt)

def crear_docx_plan(datos: dict, ruta_salida: str):
    """Crea el documento .docx del plan de práctica pedagógica"""
    doc = Document()

    # Márgenes
    for seccion in doc.sections:
        seccion.top_margin = Cm(2)
        seccion.bottom_margin = Cm(2)
        seccion.left_margin = Cm(2.5)
        seccion.right_margin = Cm(2.5)

    # Tabla principal
    tabla = doc.add_table(rows=1, cols=6)
    tabla.style = 'Table Grid'

    def celda_encabezado(celda, etiqueta, valor=""):
        print(f"etiqueta={etiqueta}, valor={valor}")  # debug
        p = celda.paragraphs[0]
        p.clear()
        run_etiqueta = p.add_run(etiqueta)
        run_etiqueta.bold = True
        run_etiqueta.font.size = Pt(10)
        if valor:
            run_valor = p.add_run(f" {valor}")
            run_valor.font.size = Pt(10)

    # Fila 1 — título
    fila = tabla.rows[0].cells
    fila[0].merge(fila[5])
    fila[0].text = ""
    p = fila[0].paragraphs[0]
    run = p.add_run("PLAN DE PRÁCTICA PEDAGÓGICA")
    run.bold = True
    run.font.size = Pt(12)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Fila 2 — centro educativo y curso lectivo
    tabla.add_row()
    fila = tabla.rows[-1].cells
    fila[0].merge(fila[3])
    fila[4].merge(fila[5])
    celda_encabezado(fila[0], "Centro educativo:", datos.get("institucion", ""))
    celda_encabezado(fila[4], "Curso lectivo:", datos.get("curso_lectivo", "2026"))

    # Fila 3 — docente y nivel
    tabla.add_row()
    fila = tabla.rows[-1].cells
    fila[0].merge(fila[3])
    fila[4].merge(fila[5])
    celda_encabezado(fila[0], "Nombre de la persona docente:", datos.get("docente", ""))
    celda_encabezado(fila[4], "Nivel:", datos.get("nivel", ""))

    # Fila 4 — especialidad, modalidad, campo detallado
    tabla.add_row()
    fila = tabla.rows[-1].cells
    fila[0].merge(fila[1])
    fila[2].merge(fila[3])
    fila[4].merge(fila[5])
    celda_encabezado(fila[0], "Especialidad:", datos.get("especialidad", ""))
    celda_encabezado(fila[2], "Modalidad:", datos.get("modalidad", ""))
    celda_encabezado(fila[4], "Campo detallado:", datos.get("campo_detallado", ""))

    # Fila 5 — subárea, unidad, tiempo
    tabla.add_row()
    fila = tabla.rows[-1].cells
    fila[0].merge(fila[1])
    fila[2].merge(fila[3])
    fila[4].merge(fila[5])
    celda_encabezado(fila[0], "Subárea:", datos.get("subarea", ""))
    celda_encabezado(fila[2], "Unidad de estudio:", datos.get("unidad", ""))
    celda_encabezado(fila[4], "Tiempo estimado:", datos.get("tiempo_total", ""))

    # Fila 6 — competencias y eje política
    tabla.add_row()
    fila = tabla.rows[-1].cells
    fila[0].merge(fila[2])
    fila[3].merge(fila[5])
    celda_encabezado(fila[0], "Competencias para el desarrollo humano:", datos.get("competencias", ""))
    celda_encabezado(fila[3], "Eje política educativa:", datos.get("eje_politica", ""))

    # Fila de cabeceras de la tabla de contenido
    # Separar las tablas
    doc.add_paragraph()
    tabla = doc.add_table(rows=0, cols=6)
    tabla.style = 'Table Grid'
    tabla.add_row()
    fila = tabla.rows[-1].cells

    # Cabeceras
    cabeceras = [
        ("Resultados de\naprendizaje", 1),
        ("Saberes esenciales", 1),
        ("Estrategias para la mediación pedagógica", 2),
        ("Evidencias de aprendizaje", 1),
        ("Tiempo estimado\n(horas)", 1),
    ]

    # Columna 0 — Resultados
    p = fila[0].paragraphs[0]
    p.clear()
    run = p.add_run("Resultados de\naprendizaje")
    run.bold = True
    run.font.size = Pt(10)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Columna 1 — Saberes
    p = fila[1].paragraphs[0]
    p.clear()
    run = p.add_run("Saberes esenciales")
    run.bold = True
    run.font.size = Pt(10)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Columna 2+3 — Estrategias (fusionadas)
    fila[2].merge(fila[3])
    p = fila[2].paragraphs[0]
    p.clear()
    run = p.add_run("Estrategias para la mediación pedagógica")
    run.bold = True
    run.font.size = Pt(10)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Columna 4 — Evidencias
    p = fila[4].paragraphs[0]
    p.clear()
    run = p.add_run("Evidencias de aprendizaje")
    run.bold = True
    run.font.size = Pt(10)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Columna 5 — Tiempo
    p = fila[5].paragraphs[0]
    p.clear()
    run = p.add_run("Tiempo estimado\n(horas)")
    run.bold = True
    run.font.size = Pt(10)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Fila de subcabeceras Docente / Estudiante
    tabla.add_row()
    fila = tabla.rows[-1].cells
    fila[0].merge(fila[1])  # vacía bajo RA y Saberes
    
    p = fila[2].paragraphs[0]
    p.clear()
    run = p.add_run("Docente")
    run.bold = True
    run.font.size = Pt(10)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = fila[3].paragraphs[0]
    p.clear()
    run = p.add_run("Estudiante")
    run.bold = True
    run.font.size = Pt(10)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    fila[4].merge(fila[5])  # vacía bajo Evidencias y Tiempo
    
    # Filas de contenido
    resultados = datos.get("resultados", [])
    for resultado in resultados:
        tabla.add_row()
        fila = tabla.rows[-1].cells

        # Resultado de aprendizaje
        fila[0].paragraphs[0].add_run(resultado.get("ra", "")).font.size = Pt(10)

        # Saberes esenciales
        fila[1].paragraphs[0].add_run(resultado.get("saberes", "")).font.size = Pt(10)

        # Estrategias docente
        fila[2].paragraphs[0].add_run(resultado.get("estrategias_docente", "")).font.size = Pt(10)

        # Estrategias estudiante
        fila[3].paragraphs[0].add_run(resultado.get("estrategias_estudiante", "")).font.size = Pt(10)

        # Evidencias
        fila[4].paragraphs[0].add_run(resultado.get("evidencias", "")).font.size = Pt(10)

        # Tiempo
        fila[5].paragraphs[0].add_run(resultado.get("horas", "")).font.size = Pt(10)
    
    # Guardar
    doc.save(ruta_salida)
    return f"Documento guardado en {ruta_salida}"

if __name__ == "__main__":
    datos_prueba = {
        "institucion": "CTP de Hojancha",
        "curso_lectivo": "2026",
        "docente": "Eidy Guevara",
        "nivel": "10°",
        "especialidad": "Informática",
        "modalidad": "Técnica",
        "campo_detallado": "Desarrollo de Software",
        "subarea": "Desarrollo Web",
        "unidad": "Unidad 1",
        "tiempo_total": "48 horas",
        "competencias": "Comunicación asertiva",
        "eje_politica": "Educar para el desarrollo sostenible",
        "resultados": [
            {
                "ra": "1. Explicar la importancia de la ingeniería del software.",
                "saberes": "• La ingeniería de software\n• La naturaleza del software\n• Definición de software",
                "estrategias_docente": "• Presenta conceptos mediante técnica expositiva dialogada usando ejemplos de apps conocidas en Guanacaste.\n• Modela comparación de modelos de proceso.",
                "estrategias_estudiante": "• Construye mapa conceptual sobre ingeniería del software.\n• Elabora cuadro comparativo de modelos.",
                "evidencias": "Conocimiento: Mapa conceptual.\nDesempeño: Exposición grupal.\nProducto: Cuadro comparativo.",
                "horas": "20"
            }
        ]
    }
    crear_docx_plan(datos_prueba, "C:\\EduForge\\test_plan2.docx")
    print("Documento creado")
